import io
import re
import datetime
import calendar
import pandas as pd
import openpyxl
from openpyxl import load_workbook
from openpyxl.styles import Font
from copy import copy
from docx import Document
import streamlit as st

def tentar_converter_numero(val):
    """Converte texto numérico em int/float nativo para o Excel reconhecer como número."""
    if pd.isna(val) or val == "" or val is None:
        return ""
    if isinstance(val, (int, float)):
        return val
    val_str = str(val).strip().replace(',', '.')
    try:
        num = float(val_str)
        return int(num) if num.is_integer() else num
    except (ValueError, TypeError):
        return str(val)

def limpar_texto_xml(texto):
    """Remove caracteres inválidos de controle ASCII que corrompem documentos Word (.docx)."""
    if pd.isna(texto) or texto is None:
        return ""
    texto_str = str(texto)
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', texto_str)

def gerar_excel_bytes(dados_exportacao):
    output = io.BytesIO()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Salvamento Remição"

    for item in dados_exportacao:
        ws.append([f"NOME: {item['nome']}"])
        ws.append([f"ORGANIZAÇÃO: {item['organiz']} | FUNÇÃO: {item['funcao']} | REMUNERAÇÃO: {item['remuneracao']} | SAÍDA: {item['saida']}"])
        ws.append([])

        pivot_df = item['pivot_df']
        if not pivot_df.empty:
            headers = ["ANO"] + list(pivot_df.columns)
            ws.append(headers)
            for idx_row, row_data in pivot_df.iterrows():
                row_vals = [tentar_converter_numero(idx_row)] + [tentar_converter_numero(v) for v in row_data.values]
                ws.append(row_vals)

        ws.append(["Total de Dias:", tentar_converter_numero(item['total_dias'])])
        ws.append([])
        ws.append([])

    wb.save(output)
    return output.getvalue()

def gerar_docx_bytes(dados_exportacao):
    output = io.BytesIO()
    doc = Document()
    doc.add_heading("Espaço de Dados para Salvamento", level=1)

    for item in dados_exportacao:
        doc.add_heading(limpar_texto_xml(f"NOME: {item['nome']}"), level=2)
        p_meta = doc.add_paragraph()
        p_meta.add_run(
            limpar_texto_xml(
                f"ORGANIZAÇÃO: {item['organiz']} | "
                f"FUNÇÃO: {item['funcao']} | "
                f"REMUNERAÇÃO: {item['remuneracao']} | "
                f"SAÍDA: {item['saida']}"
            )
        )

        pivot_df = item['pivot_df']
        if not pivot_df.empty:
            headers = ["ANO"] + list(pivot_df.columns)
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = limpar_texto_xml(h)

            for idx_row, row_data in pivot_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = limpar_texto_xml(idx_row)
                for i, val in enumerate(row_data.values):
                    row_cells[i+1].text = limpar_texto_xml(val)

        p_tot = doc.add_paragraph()
        p_tot.add_run(limpar_texto_xml(f"Total de Dias: {item['total_dias']}")).bold = True
        doc.add_paragraph()

    doc.save(output)
    return output.getvalue()

def extrair_mes_ano_do_nome(nome_arquivo):
    meses = {
        "JANEIRO": "01", "FEVEREIRO": "02", "MARÇO": "03", "MARCO": "03",
        "ABRIL": "04", "MAIO": "05", "JUNHO": "06", "JULHO": "07",
        "AGOSTO": "08", "SETEMBRO": "09", "OUTUBRO": "10",
        "NOVEMBRO": "11", "DEZEMBRO": "12"
    }
    nome_upper = str(nome_arquivo).upper()
    ano_match = re.search(r'\b(20\d{2})\b', nome_upper)
    ano = ano_match.group(1) if ano_match else None
    
    mes = None
    for nome_mes, num_mes in meses.items():
        if nome_mes in nome_upper:
            mes = num_mes
            break
            
    if mes and ano:
        return f"{mes}/{ano}"
    return "SEM MÊS/ANO"

def copiar_estilo_completo(origem, destino):
    if origem.has_style:
        destino.font = copy(origem.font)
        destino.border = copy(origem.border)
        destino.fill = copy(origem.fill)
        destino.number_format = copy(origem.number_format)
        destino.protection = copy(origem.protection)
        destino.alignment = copy(origem.alignment)

def deduplicar_colunas(colunas):
    vistos = {}
    novas_colunas = []
    for col in colunas:
        col_str = str(col).strip()
        if col_str in vistos:
            vistos[col_str] += 1
            novas_colunas.append(f"{col_str} ({vistos[col_str]})")
        else:
            vistos[col_str] = 1
            novas_colunas.append(col_str)
    return novas_colunas

def extrair_valor_limpo(df, idx, col_name):
    try:
        val = df.iloc[idx][col_name]
        if isinstance(val, pd.Series):
            val = val.iloc[0]
        if pd.isna(val):
            return None
        return val.item() if hasattr(val, 'item') else val
    except:
        return None

def converter_valor_inteligente(val_str, dtype_original):
    if val_str is None or str(val_str).strip() == "":
        return None
    val_str = str(val_str).strip()
    if pd.api.types.is_integer_dtype(dtype_original):
        try:
            return int(val_str)
        except ValueError:
            pass
    elif pd.api.types.is_float_dtype(dtype_original):
        try:
            return float(val_str.replace(',', '.'))
        except ValueError:
            pass
    try:
        return float(val_str.replace(',', '.'))
    except ValueError:
        return val_str

def formatar_datas_dataframe(df_input):
    df_out = df_input.copy()
    for col in df_out.columns:
        if pd.api.types.is_datetime64_any_dtype(df_out[col]):
            df_out[col] = df_out[col].dt.strftime('%d/%m/%Y').fillna('')
        else:
            df_out[col] = df_out[col].apply(
                lambda v: "" if pd.isna(v) else (
                    v.strftime('%d/%m/%Y') if isinstance(v, (datetime.datetime, datetime.date, pd.Timestamp))
                    else (str(v).split(' ')[0] if isinstance(v, str) and (' 00:00:00' in str(v) or 'T00:00:00' in str(v)) else v)
                )
            )
    return df_out

def calcular_pascoa(ano):
    a = ano % 19
    b = ano // 100
    c = ano % 100
    d = b // 4
    e = b % 4
    f = (b + 8) // 25
    g = (b - f + 1) // 3
    h = (19 * a + b - d - g + 15) % 30
    i = c // 4
    k = c % 4
    L = (32 + 2 * e + 2 * i - h - k) % 7
    m = (a + 11 * h + 22 * L) // 451
    mes = (h + L - 7 * m + 114) // 31
    dia = ((h + L - 7 * m + 114) % 31) + 1
    return datetime.date(ano, mes, dia)

def obter_estatisticas_mes(ano, mes):
    cal = calendar.monthcalendar(ano, mes)
    pascoa = calcular_pascoa(ano)
    feriados = [
        datetime.date(ano, 1, 1),
        pascoa - datetime.timedelta(days=47),
        pascoa - datetime.timedelta(days=2),
        datetime.date(ano, 4, 21),
        datetime.date(ano, 5, 1),
        pascoa + datetime.timedelta(days=60),
        datetime.date(ano, 9, 7),
        datetime.date(ano, 10, 12),
        datetime.date(ano, 11, 2),
        datetime.date(ano, 11, 15),
        datetime.date(ano, 11, 20),
        datetime.date(ano, 12, 25),
    ]
    feriados_mes = [f for f in feriados if f.month == mes and f.year == ano]

    dias_seg_sex_total = 0
    dias_seg_sab_total = 0
    feriados_seg_sex = 0
    feriados_seg_sab = 0
    lista_feriados_detalhes = []

    for semana in cal:
        for i in range(7):
            dia = semana[i]
            if dia != 0:
                data_atual = datetime.date(ano, mes, dia)
                wd = data_atual.weekday()
                if wd < 5:
                    dias_seg_sex_total += 1
                    dias_seg_sab_total += 1
                elif wd == 5:
                    dias_seg_sab_total += 1

                if data_atual in feriados_mes:
                    if wd < 5:
                        feriados_seg_sex += 1
                        feriados_seg_sab += 1
                        lista_feriados_detalhes.append((data_atual, "Seg a Sex"))
                    elif wd == 5:
                        feriados_seg_sab += 1
                        lista_feriados_detalhes.append((data_atual, "Sábado"))

    return {
        "seg_sex_total": dias_seg_sex_total,
        "seg_sex_feriados": feriados_seg_sex,
        "seg_sex_uteis": dias_seg_sex_total - feriados_seg_sex,
        "seg_sab_total": dias_seg_sab_total,
        "seg_sab_feriados": feriados_seg_sab,
        "seg_sab_uteis": dias_seg_sab_total - feriados_seg_sab,
        "feriados_detalhes": lista_feriados_detalhes
    }

def gerar_arquivo_atualizado_bytes(source_input, header, fila, df_original, sheet_name=None):
    wb = load_workbook(io.BytesIO(source_input) if isinstance(source_input, bytes) else source_input)
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb[wb.sheetnames[0]]
    for mod in fila:
        col_target = mod['coluna']
        valor_convertido = converter_valor_inteligente(mod['novo_valor'], df_original[col_target].dtype)
        for idx in mod['indices']:
            excel_row = idx + header + 1
            ws.cell(row=excel_row, column=df_original.columns.get_loc(col_target) + 1, value=valor_convertido)

            if col_target.strip().upper() in ["SAIDA", "SAÍDA"]:
                for col_idx in range(1, ws.max_column + 1):
                    cell = ws.cell(row=excel_row, column=col_idx)
                    current_font = cell.font
                    if current_font:
                        cell.font = Font(
                            name=current_font.name,
                            size=current_font.size,
                            bold=current_font.bold,
                            italic=current_font.italic,
                            strike=current_font.strike,
                            underline=current_font.underline,
                            color="FF0000"
                        )
                    else:
                        cell.font = Font(color="FF0000")

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()

def titulo_estilizado(subtitulo=""):
    st.markdown(
        f"<div style='text-align: center; padding: 1.5rem; background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%); color: white; border-radius: 12px; margin-bottom: 1.5rem;'><h1>⚡ SINALE WEB</h1><p>{subtitulo}</p></div>",
        unsafe_allow_html=True
    )

def obter_nome_coluna_por_letra(df, colunas_disponiveis, letra):
    mapa_letras = {
        'A': 0, 'B': 1, 'C': 2, 'D': 3, 'E': 4, 'F': 5, 'G': 6, 'H': 7,
        'I': 8, 'J': 9, 'K': 10, 'L': 11, 'M': 12, 'N': 13, 'O': 14,
        'P': 15, 'Q': 16, 'R': 17, 'S': 18, 'T': 19, 'U': 20, 'V': 21,
        'W': 22, 'X': 23, 'Y': 24, 'Z': 25
    }
    idx = mapa_letras.get(letra.upper())
    if idx is not None and idx < len(colunas_disponiveis):
        return colunas_disponiveis[idx]
    return None

def gerar_config_largura_colunas(df_subset, colunas):
    config = {}
    for col in colunas:
        if col in df_subset.columns:
            nome_coluna_upper = str(col).strip().upper()
            if nome_coluna_upper == "NOME":
                tamanho_conteudo = df_subset[col].astype(str).str.len().max() if not df_subset[col].empty else 10
                if pd.isna(tamanho_conteudo):
                    tamanho_conteudo = 10
                largura_pixels = int(tamanho_conteudo * 8) + 20
                largura_pixels = max(150, min(largura_pixels, 450))
            else:
                tamanho_titulo = len(str(col))
                largura_pixels = int(tamanho_titulo * 9) + 20
                largura_pixels = max(50, largura_pixels)
            
            config[col] = st.column_config.Column(width=largura_pixels)
    return config
